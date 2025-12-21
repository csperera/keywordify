#!/usr/bin/env python3
"""
Quick test script for Keywordify.
Creates a sample document and processes it.
"""

from docx import Document
import os

# Create a sample document
print("Creating sample lecture notes...")

doc = Document()
doc.add_heading('Machine Learning Fundamentals', 0)

doc.add_paragraph(
    'Machine learning is a subset of artificial intelligence that focuses on '
    'building systems that learn from data. The primary goal is to enable '
    'computers to learn automatically without human intervention or explicit programming.'
)

doc.add_heading('Supervised Learning', level=1)
doc.add_paragraph(
    'Supervised learning is a type of machine learning where the algorithm learns '
    'from labeled training data. The model makes predictions based on input data '
    'and is corrected when its predictions are wrong. Common examples include '
    'classification and regression tasks.'
)

doc.add_heading('Neural Networks', level=1)
doc.add_paragraph(
    'A neural network is a series of algorithms that endeavors to recognize '
    'underlying relationships in a set of data through a process that mimics '
    'the way the human brain operates. Neural networks consist of layers of '
    'interconnected nodes, where each connection has a weight that adjusts '
    'as learning proceeds.'
)

doc.add_heading('Gradient Descent', level=1)
doc.add_paragraph(
    'Gradient descent is an optimization algorithm used to minimize the cost '
    'function in machine learning models. It works by iteratively moving in '
    'the direction of steepest descent as defined by the negative of the gradient. '
    'The learning rate determines the size of the steps taken during optimization.'
)

doc.add_heading('Backpropagation', level=1)
doc.add_paragraph(
    'Backpropagation is the backbone of training neural networks. It efficiently '
    'computes gradients of the loss function with respect to the weights by using '
    'the chain rule. The algorithm propagates errors backward through the network, '
    'allowing weights to be updated via gradient descent.'
)

doc.add_heading('Regularization', level=1)
doc.add_paragraph(
    'Regularization techniques prevent overfitting by adding a penalty term to '
    'the loss function. Common methods include L1 and L2 regularization, dropout, '
    'and early stopping. These techniques help the model generalize better to '
    'unseen data by constraining the complexity of the learned model.'
)

# Save the document
os.makedirs('examples', exist_ok=True)
sample_path = 'examples/sample_lecture_notes.docx'
doc.save(sample_path)

print(f"✓ Sample document created: {sample_path}")
print("\nNow you can test with:")
print(f"  python keywordify.py {sample_path}")
