"""
Document text extraction from DOCX files.
"""

from docx import Document
from typing import List


class DocxReader:
    """Read and extract text from DOCX files."""
    
    @staticmethod
    def read_text(docx_path: str) -> str:
        """
        Extract all text from a DOCX file.
        
        Args:
            docx_path: Path to DOCX file
            
        Returns:
            Full text content with paragraphs preserved
        """
        doc = Document(docx_path)
        
        # Extract all paragraph text
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():  # Skip empty paragraphs
                paragraphs.append(para.text)
        
        # Join with double newlines to preserve paragraph structure
        return '\n\n'.join(paragraphs)
    
    @staticmethod
    def get_paragraph_count(docx_path: str) -> int:
        """Get number of paragraphs in document."""
        doc = Document(docx_path)
        return len([p for p in doc.paragraphs if p.text.strip()])
